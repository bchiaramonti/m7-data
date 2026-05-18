#!/usr/bin/env python3
"""
Gerador de Documentos Normativos M7 — baseado nos templates DOCX oficiais.

USO: Este script é usado pelas skills do plugin m7-normativos para gerar
documentos .docx 100% aderentes aos templates oficiais da M7 Investimentos.

ABORDAGEM: Em vez de construir o DOCX do zero com python-docx, este script
CLONA o template oficial correspondente e substitui o conteúdo placeholder
pelos valores reais, preservando TODA a formatação original (estilos, logos,
cabeçalhos, rodapés, tabelas, cores, fontes).

Dependências: pip install python-docx
"""

import sys
import os
import re
import copy
import json
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERRO: python-docx não instalado. Execute: pip install python-docx")
    sys.exit(1)

# Paths relativos ao diretório da skill (scripts/ está dentro da skill)
SKILL_DIR = Path(__file__).resolve().parent.parent
ASSETS_DIR = SKILL_DIR / "assets"
# Templates ficam dentro de assets/ de cada skill
TEMPLATES_DIR = ASSETS_DIR

# Mapeamento de tipo para template
TEMPLATE_MAP = {
    "POL": "TPL-POL-Template-de-Politica.docx",
    "MAN": "TPL-MAN-Template-de-Manual.docx",
    "INS": "TPL-INS-Template-de-Instrucao.docx",
    "ESP": "TPL-ESP-Template-de-Especificacao-Tecnica.docx",
}

# Mapeamento de tipo para nome por extenso
TYPE_LABEL = {
    "POL": "POLÍTICA",
    "MAN": "MANUAL",
    "INS": "INSTRUÇÃO",
    "ESP": "ESPECIFICAÇÃO TÉCNICA",
}

# Frequências de revisão por tipo
REVIEW_FREQ = {
    "POL": "Anual",
    "MAN": "Semestral",
    "INS": "Trimestral",
    "ESP": "Trimestral",
}

# Aprovadores por tipo
APPROVER_LEVEL = {
    "POL": "Diretoria",
    "MAN": "Head de área",
    "INS": "Líder do processo",
    "ESP": "Analista responsável",
}

# Cores do tema normativo M7
COLORS = {
    "heading_dark": RGBColor(0x42, 0x41, 0x35),     # #424135
    "body_text": RGBColor(0x4F, 0x4E, 0x3C),         # #4F4E3C
    "cover_title": RGBColor(0x17, 0x36, 0x5D),       # #17365D
    "cover_subtitle": RGBColor(0x4F, 0x81, 0xBD),    # #4F81BD
    "table_header_bg": "424135",
    "table_alt_row_bg": "F5F3E8",
}


def clone_template(doc_type: str) -> Document:
    """Clona o template DOCX oficial correspondente ao tipo."""
    template_name = TEMPLATE_MAP.get(doc_type)
    if not template_name:
        raise ValueError(f"Tipo inválido: {doc_type}. Use: {list(TEMPLATE_MAP.keys())}")

    template_path = TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    return Document(str(template_path))


def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Substitui texto em um parágrafo preservando a formatação dos runs."""
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # Estratégia: reconstruir os runs com o texto substituído
    new_full_text = full_text.replace(old_text, new_text)

    if not paragraph.runs:
        return False

    # Preservar formatação do primeiro run e reescrever
    first_run = paragraph.runs[0]
    first_run.text = new_full_text

    # Limpar runs restantes
    for run in paragraph.runs[1:]:
        run.text = ""

    return True


def replace_text_in_document(doc, replacements: dict):
    """Substitui placeholders em todo o documento (parágrafos, tabelas, cabeçalhos, rodapés)."""
    # Parágrafos do corpo
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            replace_text_in_paragraph(paragraph, old_text, str(new_text))

    # Tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        replace_text_in_paragraph(paragraph, old_text, str(new_text))

    # Cabeçalhos e rodapés
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    for old_text, new_text in replacements.items():
                        replace_text_in_paragraph(paragraph, old_text, str(new_text))
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    for old_text, new_text in replacements.items():
                        replace_text_in_paragraph(paragraph, old_text, str(new_text))


def set_table_header_style(table):
    """Aplica estilo M7 ao cabeçalho da tabela: fundo #424135, texto branco."""
    if not table.rows:
        return

    header_row = table.rows[0]
    for cell in header_row.cells:
        # Fundo escuro
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLORS["table_header_bg"]}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        # Texto branco e negrito
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True


def set_table_alt_rows(table):
    """Aplica linhas alternadas #F5F3E8 nas tabelas."""
    for i, row in enumerate(table.rows):
        if i == 0:
            continue  # Pular cabeçalho
        if i % 2 == 0:  # Linhas pares (0-indexed, excluindo header)
            for cell in row.cells:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{COLORS["table_alt_row_bg"]}"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers: list, rows: list, after_paragraph=None):
    """Adiciona uma tabela formatada no estilo M7."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Cabeçalho
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header

    # Dados
    for i, row_data in enumerate(rows):
        for j, value in enumerate(row_data):
            table.rows[i + 1].cells[j].text = str(value)

    # Estilizar
    set_table_header_style(table)
    set_table_alt_rows(table)

    return table


def remove_template_instructions(doc):
    """Remove a seção 'Como usar este template' do documento gerado."""
    removing = False
    paragraphs_to_remove = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text == "Como usar este template":
            removing = True
            paragraphs_to_remove.append(paragraph)
            continue

        if removing:
            # Parar ao encontrar o próximo heading (ex: "Controle do Documento")
            if paragraph.style and paragraph.style.name and paragraph.style.name.startswith("Heading"):
                removing = False
                continue
            # Também parar se o texto for um título de seção conhecido
            if text in ("Controle do Documento",):
                removing = False
                continue
            paragraphs_to_remove.append(paragraph)

    # Remover os parágrafos encontrados do XML
    for paragraph in paragraphs_to_remove:
        p_element = paragraph._element
        p_element.getparent().remove(p_element)


def replace_control_table(doc, doc_type, codigo, versao, data, metadata):
    """Substitui os campos da tabela 'Controle do Documento' por posição de linha.

    Necessário porque alguns placeholders são idênticos entre linhas (ex: '[Nome, Cargo]'
    aparece tanto em 'Elaborado por' quanto em 'Aprovado por'), impossibilitando
    substituição por string simples.
    """
    if not doc.tables:
        return

    control_table = doc.tables[0]
    # Mapa: índice da linha → valor para a coluna "Valor" (cells[1])
    row_values = {
        1: codigo,                                                              # Código
        2: versao,                                                              # Versão
        3: TYPE_LABEL[doc_type].title(),                                        # Tipo
        4: metadata.get("area_nome", metadata.get("area", "")),                 # Área
        5: data,                                                                # Data
        6: metadata.get("elaborado_por", ""),                                   # Elaborado por
        7: metadata.get("aprovado_por", ""),                                    # Aprovado por
        8: "Interno",                                                           # Classificação
        9: REVIEW_FREQ[doc_type],                                               # Revisão
        10: metadata.get("documento_superior", "N/A" if doc_type == "POL" else ""),  # Doc superior
    }

    for row_idx, value in row_values.items():
        if row_idx < len(control_table.rows):
            cell = control_table.rows[row_idx].cells[1]  # Coluna "Valor"
            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    paragraph.runs[0].text = str(value)
                    for run in paragraph.runs[1:]:
                        run.text = ""
                elif paragraph.text:
                    # Parágrafo sem runs (texto direto no XML) — reescrever
                    paragraph.text = str(value)


def replace_version_table(doc, data, metadata):
    """Substitui placeholders na tabela 'Controle de Versões' (última tabela)."""
    if len(doc.tables) < 2:
        return

    version_table = doc.tables[-1]
    if len(version_table.rows) >= 2:
        row = version_table.rows[1]  # Primeira linha de dados
        cells = row.cells
        # Coluna 1: Data
        for paragraph in cells[1].paragraphs:
            replace_text_in_paragraph(paragraph, "[DD/MM/AAAA]", data)
        # Coluna 2: Autor
        for paragraph in cells[2].paragraphs:
            replace_text_in_paragraph(paragraph, "[Autor]", metadata.get("elaborado_por", ""))


def generate_from_template(doc_type: str, metadata: dict, output_path: str):
    """
    Gera um documento normativo a partir do template oficial.

    Args:
        doc_type: POL, MAN, INS ou ESP
        metadata: dict com os campos do documento (codigo, titulo, area, versao, etc.)
        output_path: caminho do arquivo .docx de saída
    """
    doc = clone_template(doc_type)

    # Preparar data formatada
    today = datetime.now().strftime("%d/%m/%Y")

    # Construir código completo se não fornecido
    codigo = metadata.get("codigo", f"{doc_type}-{metadata.get('area', 'M7')}-{metadata.get('numero', '001')}")
    versao = metadata.get("versao", "1.0")
    titulo = metadata.get("titulo", "[Título do Documento]")
    data = metadata.get("data", today)

    # ── 1. Substituições globais (texto único no documento) ──
    replacements = {
        # Capa
        f"TEMPLATE DE {TYPE_LABEL[doc_type]}": TYPE_LABEL[doc_type],
        f"{doc_type}-[ÁREA]-[NNN]": codigo,      # com acento (real no template)
        f"{doc_type}-[AREA]-[NNN]": codigo,       # sem acento (fallback)
        "[Título do Documento]": titulo,
        "[Titulo do Documento]": titulo,
        # Versão na capa
        "Versão 1.0": f"Versão {versao}",
        "Versao 1.0": f"Versão {versao}",
        # Data na capa (formato fixo do template)
        "26/02/2026": data,
    }

    replace_text_in_document(doc, replacements)

    # ── 2. Tabela Controle do Documento (por posição, evita ambiguidade) ──
    replace_control_table(doc, doc_type, codigo, versao, data, metadata)

    # ── 3. Tabela Controle de Versões ──
    replace_version_table(doc, data, metadata)

    # ── 4. Rodapé: TPL-XXX → código do documento ──
    for section in doc.sections:
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    replace_text_in_paragraph(paragraph, f"TPL-{doc_type}", codigo)

    # ── 5. Remover seção "Como usar este template" ──
    remove_template_instructions(doc)

    # Salvar
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    print(f"✓ Documento gerado: {output}")
    print(f"  Tipo: {TYPE_LABEL[doc_type]}")
    print(f"  Código: {codigo}")
    print(f"  Template base: {TEMPLATE_MAP[doc_type]}")

    return str(output)


def main():
    parser = argparse.ArgumentParser(
        description="Gerador de Documentos Normativos M7",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  %(prog)s POL --area M7 --numero 002 --titulo "Segurança da Informação" --output ./POL-M7-002.docx
  %(prog)s MAN --area PERF --numero 001 --titulo "Operação do Funil" --doc-superior POL-M7-001 --output ./MAN-PERF-001.docx
  %(prog)s INS --area PERF --numero 001 --titulo "Fechamento Mensal" --doc-superior MAN-PERF-001 --output ./INS-PERF-001.docx
  %(prog)s ESP --area PERF --numero 001 --titulo "Cálculo de KPIs" --doc-superior MAN-PERF-001 --output ./ESP-PERF-001.docx
        """
    )

    parser.add_argument("tipo", choices=["POL", "MAN", "INS", "ESP"],
                       help="Tipo do documento normativo")
    parser.add_argument("--area", required=True,
                       help="Área (M7, PERF, INV, CRE, UNI, SEG)")
    parser.add_argument("--numero", required=True,
                       help="Número sequencial (3 dígitos, ex: 001)")
    parser.add_argument("--titulo", required=True,
                       help="Título do documento")
    parser.add_argument("--versao", default="1.0",
                       help="Versão (default: 1.0)")
    parser.add_argument("--data", default=None,
                       help="Data DD/MM/AAAA (default: hoje)")
    parser.add_argument("--elaborado-por", default="",
                       help="Autor (Nome, Cargo)")
    parser.add_argument("--aprovado-por", default="",
                       help="Aprovador (Nome, Cargo)")
    parser.add_argument("--doc-superior", default="",
                       help="Código do documento superior")
    parser.add_argument("--area-nome", default="",
                       help="Nome completo da área")
    parser.add_argument("--output", required=True,
                       help="Caminho do arquivo de saída .docx")
    parser.add_argument("--metadata-json", default=None,
                       help="JSON file com metadados adicionais")

    args = parser.parse_args()

    metadata = {
        "area": args.area,
        "numero": args.numero.zfill(3),
        "titulo": args.titulo,
        "versao": args.versao,
        "data": args.data or datetime.now().strftime("%d/%m/%Y"),
        "elaborado_por": args.elaborado_por,
        "aprovado_por": args.aprovado_por,
        "documento_superior": args.doc_superior,
        "area_nome": args.area_nome or args.area,
        "codigo": f"{args.tipo}-{args.area}-{args.numero.zfill(3)}",
    }

    # Carregar metadados adicionais de JSON se fornecido
    if args.metadata_json:
        with open(args.metadata_json) as f:
            extra = json.load(f)
            metadata.update(extra)

    generate_from_template(args.tipo, metadata, args.output)


if __name__ == "__main__":
    main()
